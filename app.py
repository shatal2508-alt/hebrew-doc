import base64
import concurrent.futures
import json

import streamlit as st
import streamlit.components.v1 as components
from anthropic import Anthropic

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import fitz  # PyMuPDF - להפיכת עמודי PDF סרוקים לתמונות לצורך OCR
except ImportError:
    fitz = None


MODEL_NAME = "claude-sonnet-5"  # Claude Sonnet 5 (claude-3-5-sonnet הוצא משימוש ב-API)
MAX_TOKENS = 4096
IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "gif", "webp"}
MAX_PARALLEL_REQUESTS = 5  # מספר קריאות API מקביליות מקסימלי (חילוץ קבצים / תמלול עמודים)

st.set_page_config(
    page_title="עוזר סיכום מסמכים",
    page_icon="📝",
    layout="centered",
)

# עיצוב מקצועי ונקי, עם RTL ממוקד רק לאזור התוכן (לא לרכיבי המערכת הפנימיים
# של Streamlit — פנייה גורפת אליהם היא מה שגרמה לאייקונים/טקסט "לקפוץ" ולהתנגש).
st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Assistant:wght@400;500;600;700&display=swap');

    html, body, [data-testid="stAppViewContainer"] {
        font-family: 'Assistant', sans-serif;
    }

    /* מסתירים מיתוג ברירת מחדל של Streamlit לצורך מראה נקי */
    #MainMenu, footer, [data-testid="stToolbar"] {
        visibility: hidden;
    }

    [data-testid="stAppViewContainer"] .main .block-container {
        direction: rtl;
        text-align: right;
        max-width: 760px;
        padding-top: 2.5rem;
        padding-bottom: 3rem;
    }

    .stTextArea textarea, .stTextInput input, .stRadio, .stFileUploader {
        direction: rtl;
        text-align: right;
    }

    h1, h2, h3 {
        font-weight: 700;
    }

    .stButton button {
        width: 100%;
        border-radius: 8px;
        font-weight: 600;
        padding: 0.6rem 1rem;
    }

    [data-testid="stFileUploaderDropzone"] {
        direction: rtl;
    }

    [data-testid="stChatInput"] textarea {
        direction: rtl;
        text-align: right;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


def init_session_state() -> None:
    """מאתחל את מפתחות ה-session_state שבהם נשמר כל המידע הזמני של הסשן.

    כל הערכים חיים אך ורק בזיכרון הסשן של Streamlit: רענון הדף או סגירת
    הלשונית פותחים חיבור/סשן חדש ומאפסים אותם לחלוטין, ללא כל שמירה קבועה.
    """
    defaults = {
        "merged_text": "",   # "תיבת האם" - הטקסט המלא והממוזג מכל הקבצים שהועלו
        "summary": "",       # הסיכום הממוקד שנוצר מתוך תיבת האם
        "qa_history": [],    # רשימת (שאלה, תשובה) של שאלות ההמשך על הסיכום
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def reset_session_state() -> None:
    st.session_state.merged_text = ""
    st.session_state.summary = ""
    st.session_state.qa_history = []


def _get_streamlit_theme_base() -> str:
    """מנסה לזהות אם ה-Theme הפעיל של Streamlit (כפי שהמשתמשת בחרה באפליקציה)
    הוא 'light' או 'dark', כדי להתאים את צבעי כפתור ההעתקה בהתאם.
    מוחזר מחרוזת ריקה אם האפיון לא זמין בגרסת Streamlit הנוכחית."""
    try:
        return st.context.theme.type
    except Exception:  # noqa: BLE001 - התכונה לא קיימת בכל גרסה, זהו fallback מכוון
        return ""


def render_copy_button(text: str, key: str, label: str = "📋 העתק טקסט") -> None:
    """מרנדר כפתור מעוצב להעתקת טקסט ללוח (clipboard) של המכשיר.

    משתמש ב-navigator.clipboard.writeText עם נפילה חזרה ל-execCommand,
    כדי שיעבוד גם בדפדפני מובייל ישנים יותר. הכפתור רץ בתוך iframe נפרד
    (components.html), ולכן צריך להתאים את הרקע/הצבעים שלו ידנית ל-Theme
    הפעיל של Streamlit — אחרת הוא מופיע כקופסה לבנה מנותקת במצב Dark Mode.
    """
    safe_text = json.dumps(text)
    button_id = f"copy-btn-{key}"

    theme_base = _get_streamlit_theme_base()
    if theme_base == "dark":
        forced_style = "color: #fafafa !important; border-color: rgba(250, 250, 250, 0.35) !important;"
    elif theme_base == "light":
        forced_style = "color: #31333F !important; border-color: rgba(49, 51, 63, 0.3) !important;"
    else:
        forced_style = ""

    html_code = f"""
    <style>
      html, body {{
        background: transparent !important;
        margin: 0;
      }}
      .copy-btn {{
        width: 100%;
        box-sizing: border-box;
        border-radius: 8px;
        font-weight: 600;
        padding: 0.5rem 1rem;
        cursor: pointer;
        font-size: 0.95rem;
        font-family: 'Assistant', sans-serif;
        border: 1px solid rgba(49, 51, 63, 0.3);
        background-color: transparent;
        color: #31333F;
      }}
      @media (prefers-color-scheme: dark) {{
        .copy-btn {{
          color: #fafafa;
          border-color: rgba(250, 250, 250, 0.35);
        }}
      }}
    </style>
    <div style="direction: rtl; text-align: right;">
      <button id="{button_id}" class="copy-btn" style="{forced_style}">{label}</button>
    </div>
    <script>
      (function() {{
        const btn = document.getElementById("{button_id}");
        const originalLabel = btn.textContent;
        btn.addEventListener("click", async function() {{
          const text = {safe_text};
          try {{
            await navigator.clipboard.writeText(text);
          }} catch (err) {{
            const textarea = document.createElement("textarea");
            textarea.value = text;
            textarea.style.position = "fixed";
            textarea.style.opacity = "0";
            document.body.appendChild(textarea);
            textarea.focus();
            textarea.select();
            try {{ document.execCommand("copy"); }} catch (e) {{}}
            document.body.removeChild(textarea);
          }}
          btn.textContent = "✅ הועתק!";
          setTimeout(function() {{ btn.textContent = originalLabel; }}, 1500);
        }});
      }})();
    </script>
    """
    components.html(html_code, height=45)


def get_api_key() -> str:
    try:
        secret_key = st.secrets["ANTHROPIC_API_KEY"]
    except Exception:  # noqa: BLE001 - st.secrets raises when no secrets file exists
        secret_key = ""

    if secret_key:
        return secret_key.strip()

    with st.expander("🔧 הגדרות: הזנת מפתח API", expanded=False):
        key_input = st.text_input(
            "מפתח API של Anthropic",
            type="password",
            help="המפתח נשמר רק לסשן הנוכחי ואינו נשלח לשום מקום מלבד Anthropic.",
        )
    return key_input.strip()


def extract_text_from_pdf(api_key: str, uploaded_file) -> str:
    if PdfReader is None:
        st.error("הספרייה pypdf אינה מותקנת. הריצו: pip install pypdf")
        st.stop()
    reader = PdfReader(uploaded_file)
    pages_text = [page.extract_text() or "" for page in reader.pages]
    text = "\n\n".join(pages_text).strip()
    if text:
        return text

    # לא נמצאה שכבת טקסט - כנראה PDF סרוק (תמונה של דף, לא טקסט אמיתי).
    # נופלים חזרה על הפיכת כל עמוד לתמונה ו"קריאתו" ע"י Claude, כמו קובץ תמונה רגיל.
    if fitz is None:
        return ""
    return _ocr_scanned_pdf(api_key, uploaded_file.getvalue())


def extract_text_from_docx(uploaded_file) -> str:
    if Document is None:
        st.error("הספרייה python-docx אינה מותקנת. הריצו: pip install python-docx")
        st.stop()
    document = Document(uploaded_file)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts).strip()


def get_image_media_type(file_name: str) -> str:
    extension = file_name.rsplit(".", 1)[-1].lower()
    media_type_map = {
        "png": "image/png",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "gif": "image/gif",
        "webp": "image/webp",
    }
    return media_type_map.get(extension, "image/png")


def _transcribe_image_bytes(api_key: str, image_bytes: bytes, media_type: str) -> str:
    """מבקש מ-Claude לתמלל/לתאר תמונה בודדת (גם עמוד PDF שהופך לתמונה)."""
    client = Anthropic(api_key=api_key)
    encoded = base64.standard_b64encode(image_bytes).decode("utf-8")
    response = client.messages.create(
        model=MODEL_NAME,
        max_tokens=MAX_TOKENS,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {"type": "base64", "media_type": media_type, "data": encoded},
                    },
                    {
                        "type": "text",
                        "text": (
                            "תמלל ותאר במדויק ובעברית את כל הטקסט והתוכן הרלוונטי "
                            "המופיע בתמונה הזו, ללא פרשנות או תוספות משלך."
                        ),
                    },
                ],
            }
        ],
    )
    return "".join(block.text for block in response.content if block.type == "text").strip()


def extract_text_from_image(api_key: str, uploaded_file) -> str:
    """מבקש מ-Claude לתמלל/לתאר את תוכן התמונה, כדי שגם קבצי תמונה יזינו
    את אותה תיבת טקסט מרכזית כמו PDF ו-Word."""
    media_type = get_image_media_type(uploaded_file.name)
    return _transcribe_image_bytes(api_key, uploaded_file.getvalue(), media_type)


def _ocr_scanned_pdf(api_key: str, pdf_bytes: bytes) -> str:
    """הופך כל עמוד ב-PDF סרוק לתמונה, ומתמלל את כל העמודים במקביל
    (במקום בזה-אחר-זה) כדי לקצר את זמן ההמתנה."""
    document = fitz.open(stream=pdf_bytes, filetype="pdf")
    zoom_matrix = fitz.Matrix(200 / 72, 200 / 72)  # ~200 DPI - רזולוציה סבירה לתמלול
    page_images = [page.get_pixmap(matrix=zoom_matrix).tobytes("png") for page in document]
    document.close()

    if not page_images:
        return ""

    with concurrent.futures.ThreadPoolExecutor(
        max_workers=min(len(page_images), MAX_PARALLEL_REQUESTS)
    ) as executor:
        page_texts = list(
            executor.map(lambda img: _transcribe_image_bytes(api_key, img, "image/png"), page_images)
        )

    return "\n\n".join(text for text in page_texts if text).strip()


def extract_text_from_upload(api_key: str, uploaded_file) -> str:
    """מחלץ טקסט מקובץ יחיד בהתאם לסיומת שלו (PDF / Word / תמונה)."""
    name = uploaded_file.name
    extension = name.rsplit(".", 1)[-1].lower() if "." in name else ""

    if extension == "pdf":
        return extract_text_from_pdf(api_key, uploaded_file)
    if extension == "docx":
        return extract_text_from_docx(uploaded_file)
    if extension in IMAGE_EXTENSIONS:
        return extract_text_from_image(api_key, uploaded_file)
    raise ValueError(f"סוג קובץ לא נתמך: .{extension}")


HEBREW_QUALITY_GUIDELINES = (
    "עליך לכתוב אך ורק בעברית תקנית, רהוטה, אקדמית וטבעית לחלוטין (רמת שפת אם גבוהה). "
    "הימנע לחלוטין מתרגום מילולי מאנגלית (Literal translation) או מניסוחים שנשמעים כמו תרגום מכונה. "
    "השתמש בפיסוק נכון, במשפטים מנוסחים היטב, בחלוקה נכונה לפסקאות, ובזרימת קריאה חלקה, "
    "ללא טעויות דקדוק או הגהה."
)


def build_summary_system_prompt(focus: str) -> str:
    return (
        "אתה עוזר כתיבה מומחה בעברית תקנית וגבוהה. "
        "המשימה שלך היא לקרוא את התוכן שיסופק לך (שעשוי להיות ממוזג ממספר מסמכים) "
        "ולכתוב ממנו סיכום בעברית רהוטה, תקנית ומדויקת, ללא שום שגיאות כתיב או ניקוד שגוי. "
        "התמקד אך ורק בהיבט הבא שביקש המשתמש, והשמט כל מידע שאינו רלוונטי אליו: "
        f"\"{focus}\". "
        "אם התוכן אינו כולל מידע רלוונטי לנושא המבוקש, ציין זאת בבירור בעברית. "
        "כתוב בסגנון ברור, קולח ומקצועי, בפסקאות מסודרות.\n\n"
        f"{HEBREW_QUALITY_GUIDELINES}"
    )


def build_followup_system_prompt() -> str:
    return (
        "אתה עוזר כתיבה מומחה בעברית תקנית וגבוהה. "
        "בהמשך תקבל את התוכן המלא של המסמכים שהועלו, את הסיכום הממוקד שכבר הוכן מהם, "
        "ולבסוף שאלת המשך של המשתמשת. "
        "ענה על השאלה בעברית ברורה ומדויקת, בהתבסס אך ורק על התוכן המלא ועל הסיכום שסופקו לך. "
        "אם התשובה לשאלה אינה נמצאת במידע שסופק, ציין זאת בבירור במקום לנחש.\n\n"
        f"{HEBREW_QUALITY_GUIDELINES}"
    )


def stream_claude_summary(api_key: str, focus: str, merged_text: str):
    """Generator שמזרים את הסיכום מ-Claude חתיכת-טקסט אחרי חתיכת-טקסט (Streaming),
    כדי שאפשר יהיה להציג אותו באתר בזמן אמת עם st.write_stream."""
    client = Anthropic(api_key=api_key)
    with client.messages.stream(
        model=MODEL_NAME,
        max_tokens=MAX_TOKENS,
        system=build_summary_system_prompt(focus),
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": f"להלן התוכן לסיכום:\n\n{merged_text}",
                        # מסמנים את בלוק המסמכים כ"בר-שמירה" (ephemeral cache) כדי שקריאות
                        # המשך שישתמשו באותו תוכן מדויק ישלמו רק על קריאת מטמון זולה,
                        # ולא על עיבוד מלא מחדש של כל המסמכים.
                        "cache_control": {"type": "ephemeral"},
                    },
                ],
            },
        ],
    ) as stream:
        yield from stream.text_stream


def stream_claude_followup(api_key: str, merged_text: str, summary: str, question: str):
    """Generator שמזרים את תשובת ההמשך מ-Claude חתיכת-טקסט אחרי חתיכת-טקסט."""
    client = Anthropic(api_key=api_key)
    cached_context = (
        f"התוכן המלא של המסמכים שהועלו:\n\n{merged_text}\n\n"
        f"---\n\nהסיכום שכבר הוכן מהתוכן:\n\n{summary}"
    )
    with client.messages.stream(
        model=MODEL_NAME,
        max_tokens=MAX_TOKENS,
        system=build_followup_system_prompt(),
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": cached_context,
                        # אותו בלוק תוכן מדויק (מסמכים + סיכום) חוזר על עצמו בכל שאלת המשך
                        # באותה שיחה, ולכן הוא מועמד מושלם למטמון: החל מהשאלה השנייה
                        # משלמים רק על קריאת מטמון זולה במקום על כל הטקסט מחדש.
                        "cache_control": {"type": "ephemeral"},
                    },
                    {
                        "type": "text",
                        "text": f"---\n\nשאלת ההמשך של המשתמשת:\n{question}",
                    },
                ],
            },
        ],
    ) as stream:
        yield from stream.text_stream


init_session_state()

st.title("📝 עוזר סיכום מסמכים")

api_key = get_api_key()

if not api_key:
    st.warning("🔑 כדי להתחיל, יש להזין מפתח API של Anthropic בתפריט 'הגדרות' שלמעלה.")
    st.stop()

focus = st.text_input("במה להתמקד בסיכום?", placeholder="לדוגמה: המסקנות העיקריות וההמלצות בלבד")

input_mode = st.radio("בחרו את אופן הקלט:", ["העלאת קבצים", "הזנת טקסט"], horizontal=True)

uploaded_files = None
manual_text = ""

if input_mode == "העלאת קבצים":
    uploaded_files = st.file_uploader(
        "העלו קובץ אחד או יותר: PDF, Word (docx) או תמונה (png/jpg/jpeg/webp)",
        type=["pdf", "docx", "png", "jpg", "jpeg", "webp"],
        accept_multiple_files=True,
    )
else:
    manual_text = st.text_area("הזינו את הטקסט לסיכום:", height=250)

col_generate, col_reset = st.columns([3, 1])
generate = col_generate.button("צור סיכום", type="primary")
reset_clicked = col_reset.button("🗑️ נקה הכל")

if reset_clicked:
    reset_session_state()
    st.rerun()

summary_freshly_streamed = False

if generate:
    if not focus.strip():
        st.error("יש לציין במה להתמקד בסיכום.")
        st.stop()

    merged_parts = []

    if input_mode == "העלאת קבצים":
        if not uploaded_files:
            st.error("יש להעלות לפחות קובץ אחד.")
            st.stop()

        with st.spinner(f"מחלץ תוכן מ-{len(uploaded_files)} קבצים (במקביל)..."):
            # מריצים את החילוץ של כל הקבצים בו-זמנית (במקום אחד-אחרי-השני),
            # כדי שזמן ההמתנה הכולל יהיה בערך כמו הקובץ הכי איטי, ולא סכום כל הקבצים.
            extraction_results: list = [None] * len(uploaded_files)
            extraction_errors: list = [None] * len(uploaded_files)
            with concurrent.futures.ThreadPoolExecutor(
                max_workers=min(len(uploaded_files), MAX_PARALLEL_REQUESTS)
            ) as executor:
                future_to_index = {
                    executor.submit(extract_text_from_upload, api_key, uploaded_file): idx
                    for idx, uploaded_file in enumerate(uploaded_files)
                }
                for future in concurrent.futures.as_completed(future_to_index):
                    idx = future_to_index[future]
                    try:
                        extraction_results[idx] = future.result()
                    except Exception as exc:  # noqa: BLE001
                        extraction_errors[idx] = exc

        first_error_idx = next((i for i, exc in enumerate(extraction_errors) if exc is not None), None)
        if first_error_idx is not None:
            st.error(
                f"אירעה שגיאה בחילוץ תוכן מהקובץ '{uploaded_files[first_error_idx].name}': "
                f"{extraction_errors[first_error_idx]}"
            )
            st.stop()

        for uploaded_file, extracted in zip(uploaded_files, extraction_results):
            if not extracted:
                st.warning(
                    f"לא נמצא תוכן הניתן לחילוץ בקובץ '{uploaded_file.name}' "
                    "(ייתכן שמדובר ב-PDF סרוק שגם התמלול האוטומטי שלו לא הצליח). הקובץ דולג."
                )
                continue

            merged_parts.append(f"### קובץ: {uploaded_file.name}\n\n{extracted}")

        if not merged_parts:
            st.error("לא נמצא תוכן הניתן לחילוץ באף אחד מהקבצים שהועלו.")
            st.stop()
    else:
        if not manual_text.strip():
            st.error("יש להזין טקסט לסיכום.")
            st.stop()
        merged_parts.append(manual_text.strip())

    merged_text = "\n\n---\n\n".join(merged_parts)

    st.divider()
    st.subheader("📄 הסיכום")
    with st.container(border=True):
        try:
            summary = st.write_stream(stream_claude_summary(api_key, focus.strip(), merged_text))
        except Exception as exc:  # noqa: BLE001
            st.error(f"אירעה שגיאה בקריאה ל-API של Anthropic: {exc}")
            st.stop()
    summary_freshly_streamed = True

    st.session_state.merged_text = merged_text
    st.session_state.summary = summary
    st.session_state.qa_history = []

if st.session_state.summary:
    if not summary_freshly_streamed:
        st.divider()
        st.subheader("📄 הסיכום")
        with st.container(border=True):
            st.markdown(st.session_state.summary)
    render_copy_button(st.session_state.summary, key="summary")
    st.download_button(
        "הורידו את הסיכום כקובץ טקסט",
        data=st.session_state.summary.encode("utf-8"),
        file_name="summary.txt",
        mime="text/plain",
    )

    st.divider()
    st.subheader("💬 שאלות המשך על הסיכום")
    st.caption("אפשר לשאול הבהרות נוספות על סמך כל התוכן שהועלה ועל סמך הסיכום שנוצר.")

    for i, (question, answer) in enumerate(st.session_state.qa_history):
        with st.chat_message("user"):
            st.markdown(question)
        with st.chat_message("assistant"):
            st.markdown(answer)
            render_copy_button(answer, key=f"qa-{i}")

    followup_question = st.chat_input("שאלו שאלה נוספת או בקשו הבהרה על הסיכום...")

    if followup_question:
        with st.chat_message("user"):
            st.markdown(followup_question)

        with st.chat_message("assistant"):
            try:
                answer = st.write_stream(
                    stream_claude_followup(
                        api_key,
                        st.session_state.merged_text,
                        st.session_state.summary,
                        followup_question,
                    )
                )
            except Exception as exc:  # noqa: BLE001
                st.error(f"אירעה שגיאה בקריאה ל-API של Anthropic: {exc}")
                st.stop()
            # render_copy_button מגיע רק אחרי ש-st.write_stream סיים לחלוטין (קריאה חוסמת
            # עד תום ההזרמה), כך שהכפתור תמיד מעתיק את הטקסט המלא ולא תשובה חלקית.
            render_copy_button(answer, key=f"qa-{len(st.session_state.qa_history)}")

        st.session_state.qa_history.append((followup_question, answer))
